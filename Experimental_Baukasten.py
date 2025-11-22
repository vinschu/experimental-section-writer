# -*- coding: utf-8 -*-
"""
author: vinschu
Modified: preview shows procedure at startup and updates immediately when a procedure is selected.
Keeps previous features (dark theme, drag & drop, undo, button color customizations, etc.)
"""
import sys
import json
import re
import time
import math
from copy import deepcopy
from pathlib import Path
from PySide6 import QtWidgets, QtGui, QtCore
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

import xml.etree.ElementTree as ET
import os
import html as html_std
import urllib.parse
import urllib.request
import urllib.error
import socket

# Standard-Datei-Pfade
if getattr(sys, 'frozen', False):
    APP_DIR = Path(sys.executable).parent
else:
    APP_DIR = Path(__file__).parent.resolve()

PROCEDURES_FILE = APP_DIR / "procedures.json"
SOLVENTS_FILE = APP_DIR / "solvents.json"
TRIVIALS_FILE = APP_DIR / "trivial_names.json"

# Dark stylesheet applied at application level.
# Main background darker, windows/dialogs and their inner areas slightly lighter and unified.
DARK_STYLESHEET = """
    QMainWindow { background: #070707; color: #ffffff; }
    QWidget { color: #ffffff; background: #171717; }               /* unified inner window color */
    QDialog, QFrame { background: #171717; color: #ffffff; border: 1px solid #2a2a2a; border-radius: 3px; }
    QMenuBar { background: #070707; color: #ffffff; }
    QMenuBar::item { background: transparent; color: #ffffff; }
    QMenu { background: #111111; color: #ffffff; }

    /* Lists & textareas: make preview same as other windows */
    QListWidget { background: #1b1b1b; border: 1px solid #2a2a2a; color: #eaeaea; }
    QTextEdit { background: #171717; border: 1px solid #2a2a2a; color: #eaeaea; }
    QLineEdit, QPlainTextEdit { background: #1a1a1a; border: 1px solid #2a2a2a; color: #ffffff; }

    /* Haupt-Button (blau) */
    QPushButton {
        background-color: #2d89ef;
        color: white;
        padding: 6px 10px;
        border-radius: 4px;
        border: 1px solid #2677d9;
    }
    QPushButton:hover {
        background-color: #3b9ef8; /* etwas heller beim Hover */
    }
    QPushButton:pressed {
        background-color: #1f6fcf; /* etwas dunkler beim Drücken */
    }

    /* Secondary-Buttons (grau, z.B. Manage, Copy Preview) */
    QPushButton[secondary="true"] {
        background-color: #6c757d;
        color: white;
        border: 1px solid #60666b;
    }
    QPushButton[secondary="true"]:hover {
        background-color: #7a8085;
    }
    QPushButton[secondary="true"]:pressed {
        background-color: #5a6165;
    }

    /* Undo (red) - objectName #undoBtn */
    QPushButton#undoBtn {
        background-color: #e74c3c;
        color: white;
        padding: 6px 10px;
        border-radius: 4px;
        border: 1px solid #c43b2a;
    }
    QPushButton#undoBtn:hover {
        background-color: #ff6b5a;
    }
    QPushButton#undoBtn:pressed {
        background-color: #c43b2a;
    }

    /* Edit Output (green) - objectName #editOutputBtn */
    QPushButton#editOutputBtn {
        background-color: #28a745;
        color: white;
        padding: 6px 10px;
        border-radius: 4px;
        border: 1px solid #1f7f34;
    }
    QPushButton#editOutputBtn:hover {
        background-color: #36c05a;
    }
    QPushButton#editOutputBtn:pressed {
        background-color: #1f7f34;
    }

    QLabel { font-weight: 600; color: #f1f1f1; }
    QToolTip { background: #ffffff; color: #000000; }
"""

def load_procedures():
    if not PROCEDURES_FILE.exists():
        sample = [
            {
                "id": "proc1",
                "name": "General esterification",
                "description": "Fischer esterification variant",
                "template": "To a stirred solution of {acid} (1.0 equiv) in {solvent} (10 mL) was added {alcohol} (1.2 equiv). The mixture was cooled to {temperature} and {catalyst} (5 mol%) was added. The reaction was stirred for {time} and then quenched with {quench}."
            }
        ]
        PROCEDURES_FILE.write_text(json.dumps(sample, indent=2, ensure_ascii=False))
    try:
        with open(PROCEDURES_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return []
    procs = []
    if isinstance(data, dict):
        if "procedures" in data and isinstance(data["procedures"], list):
            raw_list = data["procedures"]
            procs = [p for p in raw_list if isinstance(p, dict)]
        else:
            items = [(k, v) for k, v in data.items() if isinstance(v, dict)]
            if items:
                for k, v in items:
                    copy = dict(v)
                    if "id" not in copy or not copy.get("id"):
                        copy["id"] = k
                    procs.append(copy)
            else:
                if all(isinstance(v, (str, int, float, bool, list, dict, type(None))) for v in data.values()):
                    procs = [data]
    elif isinstance(data, list):
        procs = [p for p in data if isinstance(p, dict)]
    else:
        return []
    normalized = []
    for i, p in enumerate(procs):
        pid = str(p.get("id") or p.get("ID") or f"proc{i+1}")
        name = p.get("name") or p.get("title") or ""
        desc = p.get("description") or ""
        template = p.get("template") or p.get("body") or ""
        normalized.append({
            "id": pid,
            "name": str(name),
            "description": str(desc),
            "template": str(template)
        })
    try:
        raw = PROCEDURES_FILE.read_text(encoding="utf-8")
        try:
            raw_obj = json.loads(raw)
        except Exception:
            raw_obj = None
        if raw_obj != normalized:
            PROCEDURES_FILE.write_text(json.dumps(normalized, indent=2, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass
    return normalized

def save_procedures(procs):
    with open(PROCEDURES_FILE, "w", encoding="utf-8") as f:
        json.dump(procs, f, indent=2, ensure_ascii=False)

def load_solvents():
    if not SOLVENTS_FILE.exists():
        default = ["THF", "Et2O", "Hexane", "DCM", "MeOH"]
        SOLVENTS_FILE.write_text(json.dumps(default, indent=2))
    try:
        with open(SOLVENTS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, list):
                return [str(x) for x in data]
    except Exception:
        pass
    return []

def save_solvents(solvents):
    with open(SOLVENTS_FILE, "w", encoding="utf-8") as f:
        json.dump(solvents, f, indent=2, ensure_ascii=False)

def load_trivial_names():
    if not TRIVIALS_FILE.exists():
        TRIVIALS_FILE.write_text(json.dumps({}, indent=2))
    try:
        with open(TRIVIALS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, dict):
                normalized = {}
                for k, v in data.items():
                    if not isinstance(k, str):
                        continue
                    kk = remove_subscript_unicode(strip_markup_to_plain_ascii(k)).strip()
                    if kk:
                        normalized[kk] = str(v)
                return normalized
    except Exception:
        pass
    return {}

def save_trivial_names(mapping):
    try:
        with open(TRIVIALS_FILE, "w", encoding="utf-8") as f:
            json.dump(mapping, f, indent=2, ensure_ascii=False)
    except Exception:
        pass

def extract_placeholders(template_text):
    return list(dict.fromkeys(re.findall(r"\{([a-zA-Z0-9_]+)\}", template_text)))

# ---------------------------
# Helpers: subscript mapping and others
# ---------------------------

SUBSCRIPT_TRANSLATION = str.maketrans({
    "0": "₀",
    "1": "₁",
    "2": "₂",
    "3": "₃",
    "4": "₄",
    "5": "₅",
    "6": "₆",
    "7": "₇",
    "8": "₈",
    "9": "₉"
})
_UNICODE_SUB_TO_ASCII = {v: k for k, v in {
    "0": "₀",
    "1": "₁",
    "2": "₂",
    "3": "₃",
    "4": "₄",
    "5": "₅",
    "6": "₆",
    "7": "₇",
    "8": "₈",
    "9": "₉"
}.items()}
SUBSCRIPT_REVERSE_TRANSLATION = str.maketrans(_UNICODE_SUB_TO_ASCII)
UNICODE_SUBSCRIPT_DIGITS = set(_UNICODE_SUB_TO_ASCII.keys())

def subscript_digits(s: str) -> str:
    if not s:
        return s
    return s.translate(SUBSCRIPT_TRANSLATION)

def remove_subscript_unicode(s: str) -> str:
    if not s:
        return s
    return s.translate(SUBSCRIPT_REVERSE_TRANSLATION)

def normalize_commas_in_number_string(s: str) -> str:
    if not s:
        return s
    return re.sub(r'(?<=\d),(?=\d)', '.', s)

def is_zero_ml(s):
    if not s:
        return True
    raw = s.strip().lower().replace("\u00A0", " ").strip()
    if "ml" in raw:
        num = re.sub(r"[^\d\.,\-eE]", "", raw)
        num = num.replace(",", ".")
        try:
            val = float(num)
            return abs(val) < 1e-9
        except Exception:
            return False
    return False

def append_eq_if_present(s):
    if not s:
        return s
    if re.search(r"\beq\.?$", s, re.IGNORECASE):
        return s
    return f"{s} eq."

def apply_subscript_markup_legacy(s: str) -> str:
    if not s:
        return s
    s = re.sub(r'(?i)<sub>\s*([0-9]+)\s*</sub>',
               lambda m: m.group(1).translate(SUBSCRIPT_TRANSLATION),
               s)
    s = re.sub(r'_\{\s*([0-9]+)\s*\}',
               lambda m: m.group(1).translate(SUBSCRIPT_TRANSLATION),
               s)
    s = re.sub(r'_([0-9]+)',
               lambda m: m.group(1).translate(SUBSCRIPT_TRANSLATION),
               s)
    return s

_MARKUP_RE = re.compile(
    r'(?P<htmlsub><sub>.*?</sub>)'
    r'|_\{\s*([^}]+)\s*\}'
    r'|_([A-Za-z0-9])'
    r'|/\{\s*([^}]+)\s*\}'
    r'|/([A-Za-z0-9])',
    flags=re.IGNORECASE | re.DOTALL
)

def render_markup_to_html(s: str) -> str:
    if s is None:
        return ""
    out = []
    last = 0
    for m in _MARKUP_RE.finditer(s):
        start, end = m.span()
        if start > last:
            out.append(html_std.escape(s[last:start]))
        htmlsub = m.group(1)
        brace_sub = m.group(2)
        single_sub = m.group(3)
        brace_i = m.group(4)
        single_i = m.group(5)
        if htmlsub:
            inner = re.sub(r'(?i)^<sub>\s*', '', htmlsub)
            inner = re.sub(r'(?i)\s*</sub>$', '', inner)
            out.append(f"<sub>{html_std.escape(inner)}</sub>")
        elif brace_sub:
            out.append(f"<sub>{html_std.escape(brace_sub)}</sub>")
        elif single_sub:
            out.append(f"<sub>{html_std.escape(single_sub)}</sub>")
        elif brace_i:
            out.append(f"<i>{html_std.escape(brace_i)}</i>")
        elif single_i:
            out.append(f"<i>{html_std.escape(single_i)}</i>")
        else:
            out.append(html_std.escape(m.group(0)))
        last = end
    if last < len(s):
        out.append(html_std.escape(s[last:]))
    return "".join(out)

def strip_markup_to_plain_ascii(s: str) -> str:
    if s is None:
        return ""
    s = re.sub(r'(?i)<sub>\s*([^<]+?)\s*</sub>', r'\1', s)
    s = re.sub(r'(?i)<i>\s*([^<]+?)\s*</i>', r'\1', s)
    s = re.sub(r'_\{\s*([^}]+)\s*\}', r'\1', s)
    s = re.sub(r'_([A-Za-z0-9])', r'\1', s)
    s = re.sub(r'/\{\s*([^}]+)\s*\}', r'\1', s)
    s = re.sub(r'/([A-Za-z0-9])', r'\1', s)
    return s

# ---------------------------
# CDXML parsing and helpers (unchanged)
# ---------------------------

def local_name(el):
    tag = el.tag
    if isinstance(tag, str) and tag.startswith("{"):
        return tag.split("}", 1)[1]
    return tag

def extract_text_elem(elem):
    if elem is None:
        return ""
    return "".join(elem.itertext()).strip()

def parse_stoichiometry_components(cdxml_path):
    try:
        tree = ET.parse(cdxml_path)
    except ET.ParseError as e:
        raise RuntimeError(f"XML parse error for {cdxml_path}: {e}")
    root = tree.getroot()
    stoich = None
    for el in root.iter():
        if isinstance(el.tag, str) and local_name(el).lower() == "stoichiometrygrid":
            stoich = el
            break
    if stoich is None:
        return []
    components = [c for c in stoich if local_name(c).lower() == "sgcomponent"]
    if not components:
        components = [el for el in stoich.iter() if local_name(el).lower() == "sgcomponent"]
    if not components:
        return []
    header_comp = next((c for c in components if (c.attrib.get("ComponentIsHeader") or "").lower() == "yes"), None)
    headers = []
    if header_comp is not None:
        sgdata_elems = [el for el in header_comp if local_name(el).lower() == "sgdatum"]
        for sg in sgdata_elems:
            t_elem = next((d for d in sg.iter() if local_name(d).lower() == "t"), None)
            headers.append(extract_text_elem(t_elem) or "")
    else:
        first = components[0] if components else None
        if first is not None:
            sgdata_elems = [el for el in first if local_name(el).lower() == "sgdatum"]
            for sg in sgdata_elems:
                t_elem = next((d for d in sg.iter() if local_name(d).lower() == "t"), None)
                headers.append(extract_text_elem(t_elem) or "")
    data_components = [c for c in components if c is not header_comp]
    if len(data_components) > 0:
        data_components = data_components[1:]
    parsed = []
    for comp in data_components:
        sgdatums = [el for el in comp if local_name(el).lower() == "sgdatum"]
        values = []
        for sg in sgdatums:
            t_elem = next((d for d in sg.iter() if local_name(d).lower() == "t"), None)
            values.append(extract_text_elem(t_elem) or "")
        if len(values) < len(headers):
            values += [""] * (len(headers) - len(values))
        raw_fields = {}
        for i, hdr in enumerate(headers):
            key = (hdr or "").strip() or f"col_{i}"
            raw_fields[key] = values[i] if i < len(values) else ""
        formula = ""
        for k in raw_fields.keys():
            if k.strip().lower() == "formula":
                formula = raw_fields[k].strip()
                break
        if not formula and values:
            formula = values[0].strip()
        formula = re.sub(r'^[\s\-_\.0-9:()]+', '', formula).strip()
        if formula:
            parsed.append({
                "formula": formula,
                "headers": headers,
                "values": values,
                "raw_fields": raw_fields
            })
    return parsed

def find_cdxml_files(root_folder):
    cdxml_files = []
    for dirpath, _, filenames in os.walk(root_folder):
        for fname in filenames:
            if fname.lower().endswith((".cdxml", ".xml")):
                full = os.path.join(dirpath, fname)
                try:
                    with open(full, "rb") as fh:
                        head = fh.read(400).lower()
                    if b"cdxml" in head or b"chem" in head or fname.lower().endswith(".cdxml"):
                        cdxml_files.append(full)
                except Exception:
                    if fname.lower().endswith(".cdxml"):
                        cdxml_files.append(full)
    return cdxml_files

def format_component_display_by_index(component, is_last=False):
    """
    Extrahiert relevante Zeilen aus component['values'] und baut die Anzeige zusammen.

    Änderungen:
    - Für alle verwendeten Zeilen (row0, row4, row5, row9, row10, row11) wird zuerst
      normalize_commas_in_number_string angewendet und danach ein Leerzeichen zwischen
      Zahl und vielen gebräuchlichen Einheiten eingefügt (z. B. "10mg" -> "10 mg").
      Prozentangaben bleiben unverändert (z. B. "5%").
    - Verhalten für die letzte Zeile: Anzeige als "row1 (row9, row10, row11)" wie gewünscht.
    """
    import re

    vals = component.get("values", [])

    def get_last_col_value(row_index):
        try:
            if not vals:
                return ""
            first = vals[0]
            # Fall: values ist Liste von Zeilen (jede Zeile ist Liste/Tuple von Spalten)
            if isinstance(first, (list, tuple)):
                # Wenn row_index außerhalb liegt, fange Exception ab
                return (vals[row_index][-1] or "").strip() if row_index < len(vals) and vals[row_index] else ""
            # Fall: values ist Liste von Spalten (jede Spalte ist Liste/Tuple von Zeilen)
            last = vals[-1]
            if isinstance(last, (list, tuple)):
                return (last[row_index] or "").strip() if row_index < len(last) else ""
            # Fallback: flache Liste (einspaltig)
            return (vals[row_index] or "").strip() if row_index < len(vals) else ""
        except Exception:
            return ""

    def get_first_col_value(row_index):
        """Wird für formula/backoff verwendet (ähnlich zum Originalverhalten)."""
        try:
            if not vals:
                return ""
            first = vals[0]
            # Liste von Zeilen -> erste Spalte ist index 0
            if isinstance(first, (list, tuple)):
                return (vals[row_index][0] or "").strip() if row_index < len(vals) and vals[row_index] else ""
            # Liste von Spalten -> erste Spalte ist vals[0]
            if isinstance(vals[0], (list, tuple)):
                col0 = vals[0]
                return (col0[row_index] or "").strip() if row_index < len(col0) else ""
            # Fallback: flache Liste
            return (vals[row_index] or "").strip() if row_index < len(vals) else ""
        except Exception:
            return ""

    def space_before_units(s: str) -> str:
        """
        Fügt ein Leerzeichen zwischen Zahl und gebräuchlichen Einheiten ein.
        Belässt reine Prozentangaben '5%' unverändert.
        """
        if not s:
            return s
        # Entferne überflüssige Leerzeichen vor '%' (so bleibt '5%' ohne Leerzeichen)
        s = re.sub(r'\s+%', '%', s)

        # Erweiterbare Einheitensammlung (große Auswahl gebräuchlicher Einheiten)
        units = [
            r'mg', r'g', r'kg', r'µg', r'ug', r'μg',
            r'mol', r'mmol', r'µmol', r'umol', r'cmol', r'mM', r'M',
            r'mol/L', r'mol·L-1', r'mol L-1', r'mol·kg-1',
            r'L', r'l', r'mL', r'mmL', r'mml', r'ml', r'µL', r'uL', r'nL',
            r'µg/mL', r'ug/mL', r'wt%', r'v/v', r'v/v\%', r'ppm',
            r'eq\.?', r'equiv\.?', r'equiv', r'U', r'°C', r'K',
            r'bar', r'kPa', r'atm', r'mmHg', r'cm', r'mm', r'µm', r'nm'
        ]
        units_pattern = r'(?:' + r'|'.join(units) + r')\b'

        # Ersetze Muster wie "10mg", "1,0mol", "2e-3mol" -> "10 mg", "1,0 mol", "2e-3 mol"
        def repl(m):
            num = m.group('num')
            unit = m.group('unit')
            return f"{num} {unit}"

        s = re.sub(
            rf'(?i)(?P<num>[-+]?\d[\d\.,eE+\-]*)\s*(?P<unit>{units_pattern})',
            repl,
            s
        )

        return s

    def process_disp(raw: str) -> str:
        """Normalize commas and add spacing to units."""
        if not raw:
            return ""
        normalized = normalize_commas_in_number_string(raw)
        spaced = space_before_units(normalized)
        return spaced

    # Extract relevant raw values (using first/last-col heuristics)
    row0_raw = get_first_col_value(0)
    row4_raw = get_last_col_value(4)
    row5_raw = get_last_col_value(5)
    row9_raw = get_last_col_value(9)
    row10_raw = get_last_col_value(10)
    row11_raw = get_last_col_value(11)

    # Process displays for all these rows
    row0_disp = process_disp(row0_raw)
    row4_disp = process_disp(row4_raw)
    row5_disp = process_disp(row5_raw)
    row9_disp = process_disp(row9_raw)
    row10_disp = process_disp(row10_raw)
    row11_disp = process_disp(row11_raw)

    # row4_eq_disp should append "eq." if necessary (apply on processed form)
    row4_eq_disp = append_eq_if_present(row4_disp) if row4_disp else ""

    # chosen logic: prefer row9 unless it's an effective zero-ml, else row5
    chosen = row9_disp if (row9_raw and not is_zero_ml(row9_raw)) else row5_disp

    # formula fallback logic (retain original cleaning and subscript display)
    formula_raw = component.get("formula", "").strip() or row0_raw
    formula_raw = re.sub(r'^[\s\-_\.0-9:()]+', '', formula_raw).strip()
    formula_display = subscript_digits(formula_raw)

    if is_last:
        # For last row show "row1 (row9, row10, row11)" as requested.
        # For row1 we keep subscripted formula-like display (row0 may be a formula)
        row1_label = subscript_digits(row0_raw) if row0_raw else formula_display or formula_raw
        bracket_parts = []
        for part in (row9_disp, row10_disp, row11_disp):
            if part:
                bracket_parts.append(part)
        bracket = ", ".join(bracket_parts)
        display = f"{row1_label} ({bracket})" if bracket else row1_label
        return {"formula": formula_raw, "display": display, "details": {"row1": row1_label, "row9": row9_disp, "row10": row10_disp, "row11": row11_disp}}

    # Non-last: original behaviour: formula_display with bracket built from chosen, row10, row4_eq
    bracket_parts = []
    for part in (chosen, row10_disp, row4_eq_disp):
        if part:
            bracket_parts.append(part)
    bracket = ", ".join(bracket_parts)
    display = f"{formula_display} ({bracket})" if bracket else formula_display
    return {"formula": formula_raw, "display": display, "details": {"chosen": chosen, "row10": row10_disp, "row4_eq": row4_eq_disp}}
# ---------------------------
# Draggable list widget: ensure plain-text mime during drag
# ---------------------------

class DraggableListWidget(QtWidgets.QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setDragEnabled(True)
        self.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection)

    def startDrag(self, supportedActions):
        item = self.currentItem()
        if not item:
            return
        mime = QtCore.QMimeData()
        mime.setText(item.text())
        drag = QtGui.QDrag(self)
        drag.setMimeData(mime)
        pix = QtGui.QPixmap(200, 20)
        pix.fill(QtGui.QColor("#2a2a2a"))
        painter = QtGui.QPainter(pix)
        painter.setPen(QtGui.QColor("#ffffff"))
        painter.drawText(6, 14, item.text()[:60])
        painter.end()
        drag.setPixmap(pix)
        drag.exec(QtCore.Qt.CopyAction)

# ---------------------------
# Preview text widget (accepts drops from lists)
# ---------------------------

class PreviewText(QtWidgets.QTextEdit):
    dropped = QtCore.Signal(str, QtCore.QPoint)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setReadOnly(True)
        self.setAcceptDrops(True)

    def dragEnterEvent(self, event):
        md = event.mimeData()
        if md.hasText() or md.hasHtml():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        md = event.mimeData()
        if md.hasText() or md.hasHtml():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event):
        md = event.mimeData()
        if not (md.hasText() or md.hasHtml()):
            event.ignore()
            return
        text = md.text() if md.hasText() else md.html()
        pos = event.position().toPoint() if hasattr(event, 'position') else event.pos()
        self.dropped.emit(text, pos)
        event.acceptProposedAction()

# ---------------------------
# Remaining dialogs (inherit app stylesheet)
# ---------------------------

class SolventEditorDialog(QtWidgets.QDialog):
    def __init__(self, parent=None, solvent=None):
        super().__init__(parent)
        self.setWindowTitle("Solvent Editor")
        self.resize(400, 120)
        layout = QtWidgets.QVBoxLayout(self)
        form = QtWidgets.QFormLayout()
        self.name_edit = QtWidgets.QLineEdit()
        form.addRow("Solvent name:", self.name_edit)
        layout.addLayout(form)
        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Save | QtWidgets.QDialogButtonBox.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)
        if solvent:
            self.name_edit.setText(solvent)
    def get_data(self):
        return self.name_edit.text().strip()

class SolventManagerDialog(QtWidgets.QDialog):
    def __init__(self, parent=None, solvents=None):
        super().__init__(parent)
        self.setWindowTitle("Manage Solvents")
        self.resize(500, 350)
        self.solvents = list(solvents or [])
        layout = QtWidgets.QVBoxLayout(self)
        self.list_widget = QtWidgets.QListWidget()
        layout.addWidget(self.list_widget)
        self.refresh_list()
        btn_layout = QtWidgets.QHBoxLayout()
        add_btn = QtWidgets.QPushButton("Add")
        edit_btn = QtWidgets.QPushButton("Edit")
        del_btn = QtWidgets.QPushButton("Delete")
        btn_layout.addWidget(add_btn)
        btn_layout.addWidget(edit_btn)
        btn_layout.addWidget(del_btn)
        layout.addLayout(btn_layout)
        add_btn.clicked.connect(self.add_solvent)
        edit_btn.clicked.connect(self.edit_solvent)
        del_btn.clicked.connect(self.delete_solvent)
        close_btn = QtWidgets.QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)
    def refresh_list(self):
        self.list_widget.clear()
        for s in self.solvents:
            self.list_widget.addItem(subscript_digits(s))
    def add_solvent(self):
        dlg = SolventEditorDialog(self)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            val = dlg.get_data()
            if val:
                self.solvents.append(val)
                self.refresh_list()
    def edit_solvent(self):
        item = self.list_widget.currentItem()
        if not item:
            return
        idx = self.list_widget.currentRow()
        val = self.solvents[idx] if 0 <= idx < len(self.solvents) else ""
        dlg = SolventEditorDialog(self, val)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            newv = dlg.get_data()
            if newv:
                self.solvents[idx] = newv
                self.refresh_list()
    def delete_solvent(self):
        item = self.list_widget.currentItem()
        if not item:
            return
        idx = self.list_widget.currentRow()
        if idx < 0 or idx >= len(self.solvents):
            return
        confirm = QtWidgets.QMessageBox.question(self, "Delete", f"Delete solvent '{self.solvents[idx]}'?")
        if confirm == QtWidgets.QMessageBox.StandardButton.Yes:
            del self.solvents[idx]
            self.refresh_list()

class ProcedureEditorDialog(QtWidgets.QDialog):
    def __init__(self, parent=None, procedure=None):
        super().__init__(parent)
        self.setWindowTitle("Procedure Editor")
        self.resize(700, 400)
        layout = QtWidgets.QVBoxLayout(self)
        form = QtWidgets.QFormLayout()
        self.name_edit = QtWidgets.QLineEdit()
        self.desc_edit = QtWidgets.QLineEdit()
        self.template_edit = QtWidgets.QPlainTextEdit()
        form.addRow("Name:", self.name_edit)
        form.addRow("Description:", self.desc_edit)
        form.addRow("Template (use {placeholders}):", self.template_edit)
        layout.addLayout(form)
        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Save | QtWidgets.QDialogButtonBox.Cancel)
        save_btn = btns.button(QtWidgets.QDialogButtonBox.Save)
        cancel_btn = btns.button(QtWidgets.QDialogButtonBox.Cancel)
        if save_btn is not None:
            save_btn.clicked.connect(self.on_save)
        else:
            btns.accepted.connect(self.on_save)
        if cancel_btn is not None:
            cancel_btn.clicked.connect(self.reject)
        else:
            btns.rejected.connect(self.reject)
        layout.addWidget(btns)
        if procedure:
            self.name_edit.setText(procedure.get("name", ""))
            self.desc_edit.setText(procedure.get("description", ""))
            self.template_edit.setPlainText(procedure.get("template", ""))
    def on_save(self):
        self.accept()
    def get_data(self):
        return {
            "name": self.name_edit.text().strip(),
            "description": self.desc_edit.text().strip(),
            "template": self.template_edit.toPlainText().strip()
        }

class ProcedureManagerDialog(QtWidgets.QDialog):
    def __init__(self, parent=None, procedures=None):
        super().__init__(parent)
        self.setWindowTitle("Manage Procedures")
        self.resize(800, 450)
        self.procedures = [dict(p) for p in (procedures or [])]
        layout = QtWidgets.QVBoxLayout(self)
        self.list_widget = QtWidgets.QListWidget()
        layout.addWidget(self.list_widget)
        btn_row = QtWidgets.QHBoxLayout()
        add_btn = QtWidgets.QPushButton("Add")
        edit_btn = QtWidgets.QPushButton("Edit")
        del_btn = QtWidgets.QPushButton("Delete")
        up_btn = QtWidgets.QPushButton("Move Up")
        down_btn = QtWidgets.QPushButton("Move Down")
        btn_row.addWidget(add_btn)
        btn_row.addWidget(edit_btn)
        btn_row.addWidget(del_btn)
        btn_row.addWidget(up_btn)
        btn_row.addWidget(down_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)
        add_btn.clicked.connect(self.add_procedure)
        edit_btn.clicked.connect(self.edit_procedure)
        del_btn.clicked.connect(self.delete_procedure)
        up_btn.clicked.connect(self.move_up)
        down_btn.clicked.connect(self.move_down)
        bottom = QtWidgets.QHBoxLayout()
        save_btn = QtWidgets.QPushButton("Save && Close")
        cancel_btn = QtWidgets.QPushButton("Cancel")
        bottom.addStretch()
        bottom.addWidget(save_btn)
        bottom.addWidget(cancel_btn)
        layout.addLayout(bottom)
        save_btn.clicked.connect(self.accept)
        cancel_btn.clicked.connect(self.reject)
        self.refresh_list()
    def refresh_list(self):
        self.list_widget.clear()
        for p in self.procedures:
            name = p.get("name", "<no name>")
            desc = p.get("description", "")
            item = QtWidgets.QListWidgetItem(f"{name} — {desc}")
            self.list_widget.addItem(item)
    def add_procedure(self):
        dlg = ProcedureEditorDialog(self, procedure=None)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            data = dlg.get_data()
            pid = f"proc{int(time.time() * 1000)}"
            entry = {"id": pid, "name": data.get("name", ""), "description": data.get("description", ""), "template": data.get("template", "")}
            self.procedures.append(entry)
            self.refresh_list()
    def edit_procedure(self):
        item = self.list_widget.currentItem()
        if not item:
            return
        idx = self.list_widget.currentRow()
        proc = self.procedures[idx]
        dlg = ProcedureEditorDialog(self, procedure=proc)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            data = dlg.get_data()
            proc["name"] = data.get("name", proc.get("name", ""))
            proc["description"] = data.get("description", proc.get("description", ""))
            proc["template"] = data.get("template", proc.get("template", ""))
            self.procedures[idx] = proc
            self.refresh_list()
    def delete_procedure(self):
        item = self.list_widget.currentItem()
        if not item:
            return
        idx = self.list_widget.currentRow()
        proc = self.procedures[idx]
        confirm = QtWidgets.QMessageBox.question(self, "Delete", f"Delete procedure '{proc.get('name', '')}'?")
        if confirm == QtWidgets.QMessageBox.StandardButton.Yes:
            del self.procedures[idx]
            self.refresh_list()
    def move_up(self):
        idx = self.list_widget.currentRow()
        if idx <= 0:
            return
        self.procedures[idx - 1], self.procedures[idx] = self.procedures[idx], self.procedures[idx - 1]
        self.refresh_list()
        self.list_widget.setCurrentRow(idx - 1)
    def move_down(self):
        idx = self.list_widget.currentRow()
        if idx < 0 or idx >= len(self.procedures) - 1:
            return
        self.procedures[idx + 1], self.procedures[idx] = self.procedures[idx], self.procedures[idx + 1]
        self.refresh_list()
        self.list_widget.setCurrentRow(idx + 1)

class TrivialNamesManagerDialog(QtWidgets.QDialog):
    def __init__(self, parent=None, mapping=None):
        super().__init__(parent)
        self.setWindowTitle("Manage Trivial Names")
        self.resize(600, 400)
        self.mapping = dict(mapping or {})
        layout = QtWidgets.QVBoxLayout(self)
        self.list_widget = QtWidgets.QListWidget()
        layout.addWidget(self.list_widget)
        btn_layout = QtWidgets.QHBoxLayout()
        add_btn = QtWidgets.QPushButton("Add")
        edit_btn = QtWidgets.QPushButton("Edit")
        del_btn = QtWidgets.QPushButton("Delete")
        import_btn = QtWidgets.QPushButton("Import JSON...")
        btn_layout.addWidget(add_btn)
        btn_layout.addWidget(edit_btn)
        btn_layout.addWidget(del_btn)
        btn_layout.addWidget(import_btn)
        layout.addLayout(btn_layout)
        add_btn.clicked.connect(self.add_entry)
        edit_btn.clicked.connect(self.edit_entry)
        del_btn.clicked.connect(self.delete_entry)
        import_btn.clicked.connect(self.import_json)
        close_btn = QtWidgets.QPushButton("Save && Close")
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)
        self.refresh_list()
    def refresh_list(self):
        self.list_widget.clear()
        for k, v in sorted(self.mapping.items(), key=lambda x: x[0]):
            disp_key = subscript_digits(k)
            self.list_widget.addItem(f"{disp_key} → {v}")
    def add_entry(self):
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Add trivial name")
        form = QtWidgets.QFormLayout(dlg)
        formula_edit = QtWidgets.QLineEdit()
        name_edit = QtWidgets.QLineEdit()
        form.addRow("Formula (ASCII, e.g. C18H20):", formula_edit)
        form.addRow("Trivial name:", name_edit)
        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Save | QtWidgets.QDialogButtonBox.Cancel)
        btns.accepted.connect(dlg.accept)
        btns.rejected.connect(dlg.reject)
        form.addWidget(btns)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            f = formula_edit.text().strip()
            n = name_edit.text().strip()
            if f and n:
                f_norm = remove_subscript_unicode(strip_markup_to_plain_ascii(f)).strip()
                self.mapping[f_norm] = n
                self.refresh_list()
    def edit_entry(self):
        item = self.list_widget.currentItem()
        if not item:
            return
        idx = self.list_widget.currentRow()
        key = sorted(self.mapping.keys())[idx]
        val = self.mapping.get(key, "")
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Edit trivial name")
        form = QtWidgets.QFormLayout(dlg)
        formula_edit = QtWidgets.QLineEdit(key)
        name_edit = QtWidgets.QLineEdit(val)
        form.addRow("Formula (ASCII):", formula_edit)
        form.addRow("Trivial name:", name_edit)
        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Save | QtWidgets.QDialogButtonBox.Cancel)
        btns.accepted.connect(dlg.accept)
        btns.rejected.connect(dlg.reject)
        form.addWidget(btns)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            f = formula_edit.text().strip()
            n = name_edit.text().strip()
            if f:
                f_norm = remove_subscript_unicode(strip_markup_to_plain_ascii(f)).strip()
                if f_norm != key and key in self.mapping:
                    del self.mapping[key]
                self.mapping[f_norm] = n
                self.refresh_list()
    def delete_entry(self):
        item = self.list_widget.currentItem()
        if not item:
            return
        idx = self.list_widget.currentRow()
        key = sorted(self.mapping.keys())[idx]
        confirm = QtWidgets.QMessageBox.question(self, "Delete", f"Delete entry for '{key}'?")
        if confirm == QtWidgets.QMessageBox.StandardButton.Yes:
            del self.mapping[key]
            self.refresh_list()
    def import_json(self):
        fpath, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Import trivial names JSON", str(Path.cwd()), "JSON files (*.json);;All files (*)")
        if not fpath:
            return
        try:
            with open(fpath, "r", encoding="utf-8") as fh:
                data = json.load(fh)
                if isinstance(data, dict):
                    for k, v in data.items():
                        kk = remove_subscript_unicode(strip_markup_to_plain_ascii(k)).strip()
                        if kk:
                            self.mapping[kk] = str(v)
                    self.refresh_list()
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Failed to import: {e}")

class EditOutputDialog(QtWidgets.QDialog):
    def __init__(self, parent=None, initial_markup=""):
        super().__init__(parent)
        self.setWindowTitle("Edit Output (temporary)")
        self.resize(700, 400)
        layout = QtWidgets.QVBoxLayout(self)
        self.editor = QtWidgets.QPlainTextEdit()
        self.editor.setPlainText(initial_markup or "")
        # Increase editor font size by 3 points relative to preview font if available
        try:
            font = self.editor.font()
            base_pt = None
            if parent is not None and hasattr(parent, "preview_font_pt") and isinstance(parent.preview_font_pt, (int, float)):
                base_pt = int(parent.preview_font_pt)
            else:
                # fallback to application's default font point size or editor's current size
                base_pt = int(font.pointSize()) if font.pointSize() > 0 else int(QtWidgets.QApplication.font().pointSize() or 12)
            font.setPointSize(max(1, base_pt + 3))
            self.editor.setFont(font)
        except Exception:
            # non-fatal: ignore font adjustments on failure
            pass
        layout.addWidget(self.editor)
        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Save | QtWidgets.QDialogButtonBox.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)
    def get_text(self):
        return self.editor.toPlainText()

class AssignedMappingDialog(QtWidgets.QDialog):
    def __init__(self, parent=None, placeholders=None, mapping=None):
        super().__init__(parent)
        self.setWindowTitle("Assigned Mapping")
        self.resize(500, 350)
        self.placeholders = list(placeholders or [])
        self.mapping = dict(mapping or {})
        layout = QtWidgets.QVBoxLayout(self)
        self.list_widget = QtWidgets.QListWidget()
        layout.addWidget(self.list_widget)
        btn_row = QtWidgets.QHBoxLayout()
        unassign_btn = QtWidgets.QPushButton("Unassign Selected")
        copy_btn = QtWidgets.QPushButton("Copy to Clipboard")
        btn_row.addWidget(unassign_btn)
        btn_row.addWidget(copy_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)
        close_btn = QtWidgets.QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)
        unassign_btn.clicked.connect(self.unassign_selected)
        copy_btn.clicked.connect(self.copy_to_clipboard)
        self.refresh_list()
    def refresh_list(self):
        self.list_widget.clear()
        for ph in self.placeholders:
            val = self.mapping.get(ph, "")
            li = QtWidgets.QListWidgetItem(f"{ph} → {val or '<unassigned>'}")
            self.list_widget.addItem(li)
    def unassign_selected(self):
        item = self.list_widget.currentItem()
        if not item:
            return
        idx = self.list_widget.currentRow()
        ph = self.placeholders[idx]
        if ph in self.mapping:
            del self.mapping[ph]
        self.refresh_list()
    def copy_to_clipboard(self):
        lines = []
        for ph in self.placeholders:
            val = self.mapping.get(ph, "")
            lines.append(f"{ph} -> {val or ''}")
        cb = QtWidgets.QApplication.clipboard()
        cb.setText("\n".join(lines))

class DropArea(QtWidgets.QWidget):
    filesDropped = QtCore.Signal(list)
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setAutoFillBackground(True)
        pal = self.palette()
        self.setPalette(pal)
    def dragEnterEvent(self, event):
        md = event.mimeData()
        if md.hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()
    def dragMoveEvent(self, event):
        md = event.mimeData()
        if md.hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()
    def dropEvent(self, event):
        md = event.mimeData()
        if not md.hasUrls():
            event.ignore()
            return
        paths = []
        for url in md.urls():
            local = url.toLocalFile()
            if local:
                paths.append(local)
        if paths:
            self.filesDropped.emit(paths)
            event.acceptProposedAction()
        else:
            event.ignore()

class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Experimental Section Writer")
        self.resize(1100, 760)
        self.procedures = load_procedures()
        self.solvents = load_solvents()
        self.trivial_names = load_trivial_names()
        self.chemicals = []
        self.mapping = {}
        self.last_rendered_markup = ""
        self.current_template = ""
        self.preview_font_pt = None
        # history of (mapping, last_rendered_markup)
        self.state_history = []
        self.history_limit = 50
        self.setup_ui()

    def push_state(self):
        # store a deep copy of mapping and the last rendered markup
        snap = (deepcopy(self.mapping), str(self.last_rendered_markup))
        self.state_history.append(snap)
        if len(self.state_history) > self.history_limit:
            self.state_history.pop(0)

    def undo_last_action(self):
        if not self.state_history:
            QtWidgets.QMessageBox.information(self, "Undo", "Nothing to undo.")
            return
        last_map, last_markup = self.state_history.pop()
        self.mapping = deepcopy(last_map)
        self.last_rendered_markup = str(last_markup)
        # update preview based on restored markup
        self.update_preview_from_markup(self.last_rendered_markup)
        self.refresh_mapping_list()
        self.statusBar().showMessage("Undo applied", 3000)

    def setup_ui(self):
        men = self.menuBar()
        file_menu = men.addMenu("File")
        load_act = QtGui.QAction("Load CDMXL...", self)
        load_act.triggered.connect(self.load_cdmxl)
        file_menu.addAction(load_act)
        procedures_menu = men.addMenu("Procedures")
        manage_act = QtGui.QAction("Manage Procedures...", self)
        manage_act.triggered.connect(self.open_procedure_manager)
        procedures_menu.addAction(manage_act)
        solvents_act = QtGui.QAction("Solvents", self)
        solvents_act.triggered.connect(self.open_solvent_manager)
        men.addAction(solvents_act)
        names_menu = men.addMenu("Names")
        manage_triv_act = QtGui.QAction("Manage Trivial Names...", self)
        manage_triv_act.triggered.connect(self.open_trivial_names_manager)
        names_menu.addAction(manage_triv_act)
        view_menu = men.addMenu("View")
        assigned_mapping_act = QtGui.QAction("Assigned Mapping...", self)
        assigned_mapping_act.triggered.connect(self.open_assigned_mapping_dialog)
        view_menu.addAction(assigned_mapping_act)

        central = DropArea(self)
        central.filesDropped.connect(self.load_cdxml_paths)
        self.setCentralWidget(central)
        layout = QtWidgets.QHBoxLayout(central)
        layout.setContentsMargins(8, 8, 8, 8)

        left_frame = QtWidgets.QFrame()
        left_frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        left_frame.setMinimumWidth(300)
        left_layout = QtWidgets.QVBoxLayout(left_frame)
        left_layout.setSpacing(10)

        left_layout.addWidget(QtWidgets.QLabel("Chemicals (from CDXML)"))
        self.chem_list = DraggableListWidget()
        left_layout.addWidget(self.chem_list, 3)

        left_layout.addWidget(QtWidgets.QLabel("Solvents"))
        self.solv_list = DraggableListWidget()
        left_layout.addWidget(self.solv_list, 1)

        # Bottom-left Load button removed per request; keep only Manage button
        left_btn_row = QtWidgets.QHBoxLayout()
        manage_solvents_btn = QtWidgets.QPushButton("Manage")
        manage_solvents_btn.setProperty("secondary", True)
        manage_solvents_btn.clicked.connect(self.open_solvent_manager)
        left_btn_row.addWidget(manage_solvents_btn)
        left_layout.addLayout(left_btn_row)

        right_frame = QtWidgets.QFrame()
        right_frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        right_layout = QtWidgets.QVBoxLayout(right_frame)
        right_layout.setSpacing(8)

        proc_row = QtWidgets.QHBoxLayout()
        proc_row.addWidget(QtWidgets.QLabel("Procedure:"))
        self.proc_combo = QtWidgets.QComboBox()
        # Verbindung hinzufügen: bei Auswahländerung sofort on_proc_changed aufrufen
        self.proc_combo.currentIndexChanged.connect(self.on_proc_changed)
        proc_row.addWidget(self.proc_combo, 1)

        load_btn2 = QtWidgets.QPushButton("Load")
        load_btn2.setToolTip("Load CDMXL files or folders")
        load_btn2.clicked.connect(self.load_cdmxl)
        proc_row.addWidget(load_btn2)

        right_layout.addLayout(proc_row)

        right_layout.addWidget(QtWidgets.QLabel("Placeholders (click to assign)"))
        self.placeholders_list = QtWidgets.QListWidget()
        # Make placeholder area larger per user request
        self.placeholders_list.setMinimumHeight(180)
        self.placeholders_list.setMaximumHeight(280)
        self.placeholders_list.itemClicked.connect(self.on_placeholder_clicked)
        right_layout.addWidget(self.placeholders_list)

        preview_lbl = QtWidgets.QLabel("Preview")
        right_layout.addWidget(preview_lbl)
        self.preview = PreviewText()
        font = self.preview.font()
        base = font.pointSize()
        if base <= 0:
            base = QtWidgets.QApplication.font().pointSize()
            if base <= 0:
                base = 12
        font.setPointSize(base + 2)
        self.preview.setFont(font)
        self.preview_font_pt = base + 2
        # ensure objectNames for styled buttons (set later when creating buttons)
        self.preview.dropped.connect(self.on_preview_dropped)
        right_layout.addWidget(self.preview, 1)

        bottom_h = QtWidgets.QHBoxLayout()
        preview_btn = QtWidgets.QPushButton("Render Preview")
        preview_btn.clicked.connect(self.render_preview)
        bottom_h.addWidget(preview_btn)

        # Undo button (objectName used for styling)
        undo_btn = QtWidgets.QPushButton("Undo")
        undo_btn.setObjectName("undoBtn")
        undo_btn.clicked.connect(self.undo_last_action)
        bottom_h.addWidget(undo_btn)

        copy_preview_btn = QtWidgets.QPushButton("Copy Preview")
        # copy preview grey - same as Manage button (uses secondary)
        copy_preview_btn.setProperty("secondary", True)
        copy_preview_btn.clicked.connect(self.copy_preview_to_clipboard)
        bottom_h.addWidget(copy_preview_btn)

        export_word_btn = QtWidgets.QPushButton("Export to Word")
        export_word_btn.clicked.connect(self.export_word)
        bottom_h.addWidget(export_word_btn)

        edit_output_btn = QtWidgets.QPushButton("Edit Output")
        # set objectName for centralized styling
        edit_output_btn.setObjectName("editOutputBtn")
        edit_output_btn.setToolTip("Edit the generated output text (temporary).")
        edit_output_btn.clicked.connect(self.open_edit_output_dialog)
        bottom_h.addWidget(edit_output_btn)

        convert_btn = QtWidgets.QPushButton("Convert Names")
        convert_btn.setToolTip("Convert chemical formulas in the preview using trivial_names.json only (no network).")
        convert_btn.clicked.connect(self.convert_preview_formulas_to_names)
        bottom_h.addWidget(convert_btn)

        for b in (export_word_btn, edit_output_btn):
            b.setProperty("secondary", False)

        right_layout.addLayout(bottom_h)

        splitter = QtWidgets.QSplitter(QtCore.Qt.Horizontal)
        splitter.addWidget(left_frame)
        splitter.addWidget(right_frame)
        splitter.setStretchFactor(0, 0)
        splitter.setStretchFactor(1, 1)
        layout.addWidget(splitter, 1)

        status = QtWidgets.QLabel("")
        self.statusBar().addWidget(status)

        self.reload_procedures()
        self.refresh_solvents_list()

    # ---------------------------
    # Loading CDXML files (unchanged)
    # ---------------------------

    def load_cdmxl(self):
        fpath, _ = QtWidgets.QFileDialog.getOpenFileName(
            self, "Open CDMXL file (or cancel to choose folder)", str(Path.cwd()),
            "CDMXL Files (*.cdxml *.xml);;All files (*)"
        )
        if fpath:
            self.load_cdxml_paths([fpath])
            return
        return

    def load_cdxml_paths(self, paths):
    # Wenn neue Dateien/Folders geladen werden: vorherige Zuordnungen rücksetzen
    # und Preview auf ungefüllten Template-Zustand bringen.
        try:
            # Falls es aktuelle Zuordnungen oder gerenderten Text gibt, in die History pushen
            # damit das Löschen rückgängig gemacht werden kann.
            try:
                if getattr(self, "mapping", None) or getattr(self, "last_rendered_markup", None):
                    self.push_state()
            except Exception:
                # im Fehlerfall trotzdem fortfahren
                pass

            # Entferne alle aktuellen Assignments (keine Undo-Pushes - bereits oben gepusht)
            self.mapping = {}
            self.last_rendered_markup = ""
            # Clear preview widget content
            if hasattr(self, "preview") and self.preview is not None:
                try:
                    self.preview.clear()
                except Exception:
                    pass
            # Clear chemicals list UI while we load new ones
            if hasattr(self, "chem_list") and self.chem_list is not None:
                try:
                    self.chem_list.clear()
                except Exception:
                    pass
            # Clear placeholders list UI (will be repopulated if procedure selected)
            if hasattr(self, "placeholders_list") and self.placeholders_list is not None:
                try:
                    self.placeholders_list.clear()
                except Exception:
                    pass
            # Update status of mapping/UI
            try:
                self.refresh_mapping_list()
            except Exception:
                pass
        except Exception:
            # be robust: don't block loading if clearing fails
            pass

        to_parse = []
        for p in paths:
            if os.path.isfile(p):
                to_parse.append(p)
            elif os.path.isdir(p):
                found = find_cdxml_files(p)
                if found:
                    to_parse.extend(found)
        to_parse = list(dict.fromkeys(to_parse))
        if not to_parse:
            QtWidgets.QMessageBox.information(self, "No files", "No CDXML/XML files found in the dropped items.")
            return
        components_all = []
        for file in to_parse:
            try:
                comps = parse_stoichiometry_components(file)
            except Exception:
                comps = []
            if comps:
                components_all.extend(comps)
        candidates = []
        for idx, comp in enumerate(components_all):
            is_last = (idx == len(components_all) - 1)
            formatted = format_component_display_by_index(comp, is_last=is_last)
            if not formatted["formula"]:
                continue
            candidates.append({
                "formula": formatted["formula"],
                "display": formatted["display"]
            })
        seen = set()
        normalized = []
        for c in candidates:
            f = c["formula"].strip()
            if not f:
                continue
            key = f.lower()
            if key in seen:
                continue
            seen.add(key)
            normalized.append(c)
        if not normalized:
            QtWidgets.QMessageBox.information(self, "No chemicals", "No candidate chemicals found in the selected file(s).")
            # preview is already cleared at the top; ensure UI mapping state consistent
            try:
                self.refresh_mapping_list()
            except Exception:
                pass
            return
        self.chemicals = normalized
        # populate chemicals list UI
        self.chem_list.clear()
        for c in self.chemicals:
            self.chem_list.addItem(c["display"])
        self.statusBar().showMessage(f"Loaded {len(self.chemicals)} candidate chemicals", 5000)

        # Ensure assignments remain cleared and preview shows the template with placeholders unfilled.
        # (If a procedure/template is selected, render it with placeholders unassigned)
        try:
            self.mapping = {}
            # render the current template (will show placeholders as {name} if unassigned)
            self.render_preview()
            self.refresh_mapping_list()
        except Exception:
            # non-fatal: ignore rendering errors here
            pass

    def refresh_solvents_list(self):
        self.solv_list.clear()
        for s in self.solvents:
            self.solv_list.addItem(subscript_digits(s))

    def open_solvent_manager(self):
        dlg = SolventManagerDialog(self, solvents=self.solvents)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            self.solvents = dlg.solvents
            save_solvents(self.solvents)
            self.refresh_solvents_list()

    def reload_procedures(self):
        self.proc_combo.clear()
        for p in self.procedures:
            name = p.get("name") or "<no name>"
            self.proc_combo.addItem(name, p)
        if hasattr(self, "placeholders_list") and hasattr(self, "proc_combo") and self.procedures:
            self.proc_combo.setCurrentIndex(0)
            # On startup select first and render preview (on_proc_changed will call render)
            self.on_proc_changed(0)

    def open_procedure_manager(self):
        dlg = ProcedureManagerDialog(self, procedures=self.procedures)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            self.procedures = dlg.procedures
            save_procedures(self.procedures)
            self.reload_procedures()

    def on_proc_changed(self, idx):
        p = self.proc_combo.currentData()
        if not p:
            if hasattr(self, "placeholders_list"):
                self.placeholders_list.clear()
            return
        self.current_template = p.get("template", "")
        phs = extract_placeholders(self.current_template)
        self.placeholders = phs
        self.mapping = {}
        if hasattr(self, "placeholders_list"):
            self.placeholders_list.clear()
            for ph in phs:
                item = QtWidgets.QListWidgetItem(ph)
                self.placeholders_list.addItem(item)
        # render preview immediately
        self.render_preview()

    def on_placeholder_clicked(self, item):
        placeholder = item.text()

        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle(f"Assign chemical to {{{placeholder}}}")
        dlg.resize(400, 300)
        layout = QtWidgets.QVBoxLayout(dlg)

        listw = QtWidgets.QListWidget()
        layout.addWidget(listw)

        # UNASSIGN option
        listw.addItem("<unassigned>")

        # Fill list depending on placeholder type
        if "solvent" in placeholder.lower():
            # solvents first
            for s in self.solvents:
                listw.addItem(subscript_digits(s))

            # separator
            sep = QtWidgets.QListWidgetItem("── Chemicals ──")
            sep.setFlags(QtCore.Qt.NoItemFlags)
            listw.addItem(sep)

            for c in self.chemicals:
                listw.addItem(c.get("display"))
        else:
            # chemicals first
            for c in self.chemicals:
                listw.addItem(c.get("display"))

            sep = QtWidgets.QListWidgetItem("── Solvents ──")
            sep.setFlags(QtCore.Qt.NoItemFlags)
            listw.addItem(sep)

            for s in self.solvents:
                listw.addItem(subscript_digits(s))

        # Click behavior — immediate selection + close dialog
        def on_item_clicked(clicked_item):
            text = clicked_item.text()

            # push current state for undo
            self.push_state()

            if text == "<unassigned>":
                if placeholder in self.mapping:
                    del self.mapping[placeholder]
            elif "──" not in text:  # ignore separators
                self.mapping[placeholder] = text

            self.render_preview()
            dlg.accept()  # close instantly

        listw.itemClicked.connect(on_item_clicked)

        dlg.exec()

    def refresh_mapping_list(self):
        assigned = sum(1 for v in self.mapping.values() if v)
        self.statusBar().showMessage(f"{assigned}/{len(getattr(self, 'placeholders', []) or [])} placeholders assigned", 4000)

    def render_preview(self):
        if not hasattr(self, "current_template"):
            return
        text = self.current_template
        def repl(m):
            key = m.group(1)
            return self.mapping.get(key, f"{{{key}}}")
        rendered = re.sub(r"\{([a-zA-Z0-9_]+)\}", repl, text)
        self.last_rendered_markup = rendered
        self.update_preview_from_markup(rendered)
        self.refresh_mapping_list()

    def update_preview_from_markup(self, markup):
        paragraphs = []
        font_pt = self.preview_font_pt or 14
        for line in markup.splitlines() or [""]:
            frag = render_markup_to_html(line)
            paragraphs.append(f'<p style="font-size:{font_pt}pt; text-align:justify; line-height:1.5; margin:0 0 6px 0;">{frag}</p>')
        html_full = "<html><body>" + "".join(paragraphs) + "</body></html>"
        if hasattr(self, "preview"):
            self.preview.setHtml(html_full)

    def open_edit_output_dialog(self):
        if not getattr(self, "last_rendered_markup", "").strip():
            self.render_preview()
        initial = getattr(self, "last_rendered_markup", "") or ""
        dlg = EditOutputDialog(self, initial_markup=initial)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            # push previous state for undo
            self.push_state()
            new_text = dlg.get_text() or ""
            self.last_rendered_markup = new_text
            self.update_preview_from_markup(new_text)

    def _convert_unicode_subs_to_html_subs(self, s: str) -> str:
        if not s:
            return s
        chars = "".join(re.escape(ch) for ch in UNICODE_SUBSCRIPT_DIGITS)
        pattern = re.compile(f"[{chars}]+")
        def repl(m):
            uni = m.group(0)
            ascii_digits = uni.translate(SUBSCRIPT_REVERSE_TRANSLATION)
            return f"<sub>{ascii_digits}</sub>"
        return pattern.sub(repl, s)

    def copy_preview_to_clipboard(self):
        markup = getattr(self, "last_rendered_markup", "") or ""
        if not markup.strip():
            return
        markup_for_html = self._convert_unicode_subs_to_html_subs(markup)
        paragraphs = []
        font_pt = self.preview_font_pt or 14
        for line in markup_for_html.splitlines() or [""]:
            frag = render_markup_to_html(line)
            paragraphs.append(f'<p style="font-size:{font_pt}pt; text-align:justify; line-height:1.5; margin:0 0 6px 0;">{frag}</p>')
        html_full = f"<html><body>{''.join(paragraphs)}</body></html>"
        plain_lines = []
        for line in markup.splitlines() or [""]:
            ascii_line = remove_subscript_unicode(line)
            ascii_line = strip_markup_to_plain_ascii(ascii_line)
            plain_lines.append(ascii_line)
        plain_ascii = "\n".join(plain_lines)
        mime = QtCore.QMimeData()
        mime.setHtml(html_full)
        mime.setText(plain_ascii)
        cb = QtWidgets.QApplication.clipboard()
        cb.setMimeData(mime)

    def export_word(self):
        markup = getattr(self, "last_rendered_markup", "") or ""
        if not markup.strip():
            QtWidgets.QMessageBox.warning(self, "Empty", "Nothing to export")
            return
        save_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Word Document", str(Path.cwd() / "experimental_section.docx"), "Word Document (*.docx)")
        if not save_path:
            return
        token_re = re.compile(
            r'(?P<htmlsub><sub>\s*([^<]+?)\s*</sub>)'
            r'|_(?P<brace_sub>\{([^}]+)\})'
            r'|_(?P<single_sub>[A-Za-z0-9])'
            r'|/(?P<brace_i>\{([^}]+)\})'
            r'|/(?P<single_i>[A-Za-z0-9])',
            flags=re.IGNORECASE | re.DOTALL
        )
        try:
            doc = Document()
            for line in markup.splitlines() or [""]:
                p = doc.add_paragraph()
                pf = p.paragraph_format
                try:
                    pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                except Exception:
                    pass
                try:
                    pf.line_spacing = 1.5
                except Exception:
                    try:
                        pf.line_spacing = Pt(18)
                    except Exception:
                        pass
                last = 0
                for m in token_re.finditer(line):
                    if m.start() > last:
                        txt = line[last:m.start()]
                        p.add_run(txt)
                    if m.group('htmlsub'):
                        inner = re.sub(r'(?i)^<sub>\s*', '', m.group('htmlsub'))
                        inner = re.sub(r'(?i)\s*</sub>$', '', inner)
                        run = p.add_run(inner)
                        try:
                            run.font.subscript = True
                        except Exception:
                            pass
                    elif m.group('brace_sub'):
                        inner = m.group(3)
                        run = p.add_run(inner)
                        try:
                            run.font.subscript = True
                        except Exception:
                            pass
                    elif m.group('single_sub'):
                        inner = m.group('single_sub')
                        run = p.add_run(inner)
                        try:
                            run.font.subscript = True
                        except Exception:
                            pass
                    elif m.group('brace_i'):
                        inner = m.group(6)
                        run = p.add_run(inner)
                        try:
                            run.font.italic = True
                        except Exception:
                            pass
                    elif m.group('single_i'):
                        inner = m.group('single_i')
                        run = p.add_run(inner)
                        try:
                            run.font.italic = True
                        except Exception:
                            pass
                    last = m.end()
                if last < len(line):
                    p.add_run(line[last:])
            doc.save(save_path)
            QtWidgets.QMessageBox.information(self, "Saved", f"Saved to {save_path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Failed to save Word: {e}")

    def open_trivial_names_manager(self):
        dlg = TrivialNamesManagerDialog(self, mapping=self.trivial_names)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            self.trivial_names = dlg.mapping
            save_trivial_names(self.trivial_names)
            self.statusBar().showMessage(f"Saved {len(self.trivial_names)} trivial name entries", 4000)

    def open_assigned_mapping_dialog(self):
        dlg = AssignedMappingDialog(self, placeholders=getattr(self, "placeholders", []), mapping=self.mapping)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            # push current state for undo, then copy returned mapping
            self.push_state()
            self.mapping = dlg.mapping
            self.render_preview()

    def convert_preview_formulas_to_names(self):
        """
        Convert formulas in self.last_rendered_markup using only trivial_names mapping.
        No network lookups.
        """
        markup = getattr(self, "last_rendered_markup", "") or ""
        if not markup.strip():
            QtWidgets.QMessageBox.information(self, "Nothing", "No rendered text to convert.")
            return

        trivial_map = {k: v for k, v in self.trivial_names.items()}

        # find candidate tokens in markup that look like formulas:
        token_candidates = set(re.findall(r'[A-Za-z₀₁₂₃₄₅₆₇₈₉0-9\(\)]+', markup))
        candidates = []
        for tok in token_candidates:
            plain = remove_subscript_unicode(strip_markup_to_plain_ascii(tok)).strip()
            if not plain:
                continue
            if re.search(r'[A-Za-z]', plain) and re.search(r'\d', plain):
                candidates.append((tok, plain))

        if not candidates:
            QtWidgets.QMessageBox.information(self, "No formulas", "No formula-like tokens found to convert.")
            return

        # Sort to avoid partial replacements
        candidates.sort(key=lambda x: len(x[1]), reverse=True)

        new_markup = markup
        looked_up = {}

        # push previous state for undo
        self.push_state()

        for orig_tok, plain in candidates:
            plain_norm = plain.replace(" ", "")
            trivial_name = trivial_map.get(plain_norm)
            if trivial_name:
                new_markup = new_markup.replace(orig_tok, trivial_name)
                looked_up[plain_norm] = ("trivial", trivial_name)
            else:
                looked_up[plain_norm] = ("notfound", None)

        any_changed = new_markup != markup
        if any_changed:
            self.last_rendered_markup = new_markup
            self.update_preview_from_markup(new_markup)
            found_count = sum(1 for v in looked_up.values() if v[0] == "trivial")
            self.statusBar().showMessage(f"Converted {found_count} formula(s) using trivial names", 6000)
            QtWidgets.QMessageBox.information(self, "Converted", f"Converted {found_count} formula(s) in the preview.")
        else:
            QtWidgets.QMessageBox.information(self, "No changes", "No conversions were possible with trivial names.")

    # ---------------------------
    # Nearest-placeholder drop resolution (unchanged behaviour)
    # ---------------------------

    def on_preview_dropped(self, text: str, pos: QtCore.QPoint):
        if not getattr(self, "current_template", ""):
            QtWidgets.QMessageBox.information(self, "No template", "No procedure/template loaded.")
            return

        doc = self.preview.document()
        best_ph = None
        best_dist = float('inf')

        for ph in getattr(self, "placeholders", []) or []:
            visible = self.mapping.get(ph, f"{{{ph}}}")
            if not visible:
                visible = f"{{{ph}}}"
            search_text = remove_subscript_unicode(strip_markup_to_plain_ascii(visible))
            if not search_text:
                search_text = visible
            cursor = doc.find(search_text, 0)
            while cursor and not cursor.isNull():
                rect = self.preview.cursorRect(cursor)
                center = rect.center()
                dx = center.x() - pos.x()
                dy = center.y() - pos.y()
                dist = math.hypot(dx, dy)
                if dist < best_dist:
                    best_dist = dist
                    best_ph = ph
                start_after = cursor.position() + 1
                cursor = doc.find(search_text, start_after)

        if not best_ph:
            all_phs = extract_placeholders(self.current_template)
            if all_phs:
                best_ph = all_phs[0]

        if not best_ph:
            QtWidgets.QMessageBox.information(self, "No placeholders", "No placeholder was found in the template to assign to.")
            return

        val = text.strip()
        if val:
            if "──" in val:
                return
            # push state for undo
            self.push_state()
            if val == "<unassigned>":
                if best_ph in self.mapping:
                    del self.mapping[best_ph]
            else:
                self.mapping[best_ph] = val
            self.render_preview()
            self.statusBar().showMessage(f"Assigned '{val}' to placeholder {{{best_ph}}}", 3000)

def main():
    app = QtWidgets.QApplication(sys.argv)
    app.setStyleSheet(DARK_STYLESHEET)
    win = MainWindow()
    win.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()